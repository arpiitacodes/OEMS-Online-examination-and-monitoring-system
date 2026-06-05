#!/usr/bin/env python3
"""
build_docx.py — Convert the OEMS Markdown reports into professionally
formatted Microsoft Word (.docx) documents.

Handles the specific Markdown dialect used in PROJECT_REPORT.md and
VIVA_READY_QA.md: a leading title block, ATX headings (# .. ####), GitHub
pipe tables, fenced ``` code blocks, blockquotes, bold (**...**) and inline
`code`, '---' horizontal rules, and '-' bullet lists. It also inserts a real
Word Table-of-Contents field (with page numbers), a cover page, page numbers
in the footer, and consistent heading styles.

Run:  python docs/build_docx.py
"""

import os
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))

# ── Palette (matches the app's indigo/slate theme) ──
INDIGO = RGBColor(0x4F, 0x46, 0xE5)
SLATE = RGBColor(0x1E, 0x29, 0x3B)
GREY = RGBColor(0x37, 0x41, 0x51)
MUTED = RGBColor(0x6B, 0x72, 0x80)
CODE_BG = "F1F5F9"
HEADER_BG = "4F46E5"
ZEBRA_BG = "F8FAFC"


# ════════════════════════════════════════════════════════════════════
# Low-level OOXML helpers
# ════════════════════════════════════════════════════════════════════
def _shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        m.append(node)
    tcPr.append(m)


def _add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    run = p.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED

    def _field(instr):
        r = p.add_run()
        r.font.size = Pt(8)
        r.font.color.rgb = MUTED
        fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
        it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
        fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
        r._r.append(fb); r._r.append(it); r._r.append(fe)

    _field("PAGE")
    run = p.add_run(" of ")
    run.font.size = Pt(8); run.font.color.rgb = MUTED
    _field("NUMPAGES")


def _add_footer_text(section, text):
    footer = section.footer
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(7.5)
    r.font.color.rgb = MUTED


def _add_toc(doc):
    """Insert a Word TOC field (levels 1-3). Word populates it on open /
    F9; we also set a fallback message."""
    p = doc.add_paragraph()
    run = p.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = r'TOC \o "1-3" \h \z \u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    msg = OxmlElement("w:t")
    msg.text = "Right-click and choose 'Update Field' to build the Table of Contents."
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    run._r.append(fb); run._r.append(it); run._r.append(sep)
    run._r.append(msg); run._r.append(fe)


# ════════════════════════════════════════════════════════════════════
# Document styling
# ════════════════════════════════════════════════════════════════════
def _style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = GREY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    heading_specs = {
        "Heading 1": (17, INDIGO, True, 14, 6),
        "Heading 2": (13.5, SLATE, True, 12, 4),
        "Heading 3": (11.5, INDIGO, True, 8, 3),
        "Heading 4": (10.5, SLATE, True, 6, 2),
    }
    for name, (size, color, bold, before, after) in heading_specs.items():
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = bold
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def _add_runs_with_inline(paragraph, text):
    """Render **bold**, *italic*, `code`, and plain text into a paragraph.

    Tokenizing on inline-code FIRST means the *.pt / *.onnx style patterns that
    live inside `code` spans are never mistaken for italic emphasis.
    """
    # First isolate inline-code spans so their contents are treated literally.
    code_split = re.split(r"(`[^`]+`)", text)
    for chunk in code_split:
        if not chunk:
            continue
        if chunk.startswith("`") and chunk.endswith("`"):
            r = paragraph.add_run(chunk[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = INDIGO
            continue
        # In non-code text, split on bold then italic.
        for tok in re.split(r"(\*\*.+?\*\*|\*[^*\n]+?\*)", chunk):
            if not tok:
                continue
            if tok.startswith("**") and tok.endswith("**"):
                r = paragraph.add_run(tok[2:-2])
                r.bold = True
            elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
                r = paragraph.add_run(tok[1:-1])
                r.italic = True
            else:
                paragraph.add_run(tok)


# ════════════════════════════════════════════════════════════════════
# Markdown block parsing
# ════════════════════════════════════════════════════════════════════
def _parse_table(lines):
    """lines: list of pipe-table rows (header, separator, body...)."""
    rows = []
    for ln in lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    header = rows[0]
    body = rows[2:]  # skip the |---|---| separator
    return header, body


def _add_table(doc, header, body):
    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(header):
        _shade_cell(hdr_cells[i], HEADER_BG)
        _set_cell_margins(hdr_cells[i])
        cell_p = hdr_cells[i].paragraphs[0]
        cell_p.paragraph_format.space_after = Pt(0)
        _add_runs_with_inline(cell_p, htext)
        for run in cell_p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9.5)

    for ridx, row in enumerate(body):
        cells = table.add_row().cells
        for i in range(ncols):
            txt = row[i] if i < len(row) else ""
            if ridx % 2 == 1:
                _shade_cell(cells[i], ZEBRA_BG)
            _set_cell_margins(cells[i])
            cp = cells[i].paragraphs[0]
            cp.paragraph_format.space_after = Pt(0)
            _add_runs_with_inline(cp, txt)
            for run in cp.runs:
                run.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_code_block(doc, code_lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _shade_cell(cell, CODE_BG)
    _set_cell_margins(cell, top=80, bottom=80, left=140, right=140)
    cell.paragraphs[0].text = ""
    first = True
    for ln in code_lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(ln if ln else "")
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        r.font.color.rgb = SLATE
    # subtle border
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), "CBD5E1")
        tblBorders.append(b)
    table._tbl.tblPr.append(tblBorders)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    _add_runs_with_inline(p, text)
    for run in p.runs:
        run.font.italic = True
        run.font.color.rgb = MUTED
        run.font.size = Pt(9.5)
    # left accent bar
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8"); left.set(qn("w:color"), "4F46E5")
    pbdr.append(left)
    pPr.append(pbdr)


def _add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + 0.25 * level)
    p.paragraph_format.space_after = Pt(3)
    _add_runs_with_inline(p, text)


def _add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "E2E8F0")
    pbdr.append(bottom)
    pPr.append(pbdr)


# ════════════════════════════════════════════════════════════════════
# Cover page
# ════════════════════════════════════════════════════════════════════
def _add_cover(doc, title, subtitle, meta_rows):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = SLATE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(15); r.font.color.rgb = INDIGO

    # accent rule
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "4F46E5")
    pbdr.append(bottom); pPr.append(pbdr)

    for _ in range(2):
        doc.add_paragraph()

    # meta table (centered, two columns)
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in meta_rows:
        cells = table.add_row().cells
        _set_cell_margins(cells[0]); _set_cell_margins(cells[1])
        _shade_cell(cells[0], ZEBRA_BG)
        kp = cells[0].paragraphs[0]; kp.paragraph_format.space_after = Pt(0)
        kr = kp.add_run(k); kr.font.bold = True; kr.font.size = Pt(10); kr.font.color.rgb = SLATE
        vp = cells[1].paragraphs[0]; vp.paragraph_format.space_after = Pt(0)
        _add_runs_with_inline(vp, v)
        for run in vp.runs:
            run.font.size = Pt(10)
    doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# Main conversion
# ════════════════════════════════════════════════════════════════════
def convert(md_path, docx_path, cover_subtitle, footer_text):
    with open(md_path, "r", encoding="utf-8") as fh:
        raw = fh.read()
    lines = raw.split("\n")

    doc = Document()
    _style_document(doc)

    # margins
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ── Extract title block: first H1 + following '## subtitle' + the
    #    leading key/value table for the cover, then process the rest. ──
    i = 0
    # Title (first '# ')
    doc_title = "OEMS"
    while i < len(lines):
        if lines[i].startswith("# "):
            doc_title = lines[i][2:].strip()
            i += 1
            break
        i += 1
    # Optional immediate '## ' subtitle line
    cover_sub = cover_subtitle
    # gather meta rows from the first pipe-table we encounter before the first '---'
    meta_rows = []
    cover_meta_consumed_until = i
    # find first table block in the header region (before first standalone '---' rule after content)
    scan = i
    found_meta = False
    while scan < len(lines):
        ln = lines[scan]
        if ln.startswith("## ") and not cover_sub:
            cover_sub = ln[3:].strip()
        if ln.strip().startswith("|") and "|" in ln:
            # collect contiguous table
            tbl = []
            while scan < len(lines) and lines[scan].strip().startswith("|"):
                tbl.append(lines[scan]); scan += 1
            hdr, body = _parse_table(tbl)
            for row in body:
                if len(row) >= 2:
                    meta_rows.append((row[0].replace("**", ""), row[1]))
            cover_meta_consumed_until = scan
            found_meta = True
            break
        if ln.strip() == "---" and found_meta:
            break
        scan += 1

    _add_cover(doc, doc_title, cover_sub or cover_subtitle, meta_rows)

    # ── Table of Contents page ──
    doc.add_heading("Table of Contents", level=1)
    _add_toc(doc)
    doc.add_page_break()

    # ── Body: process everything AFTER the consumed cover meta table ──
    body_start = cover_meta_consumed_until if found_meta else i
    idx = body_start
    n = len(lines)

    while idx < n:
        line = lines[idx]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```"):
            idx += 1
            code = []
            while idx < n and not lines[idx].strip().startswith("```"):
                code.append(lines[idx])
                idx += 1
            idx += 1  # skip closing fence
            _add_code_block(doc, code)
            continue

        # Pipe table
        if stripped.startswith("|") and "|" in stripped:
            tbl = []
            while idx < n and lines[idx].strip().startswith("|"):
                tbl.append(lines[idx])
                idx += 1
            if len(tbl) >= 2:
                hdr, bdy = _parse_table(tbl)
                _add_table(doc, hdr, bdy)
            continue

        # Headings.
        # The source documents use '#' only for the cover title (already
        # consumed) and '##/###/####' for the actual section hierarchy. Promote
        # one level so numbered top-level sections become Word Heading 1, giving
        # the TOC a clean 1-2-3 structure: ## -> H1, ### -> H2, #### -> H3.
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            md_level = len(m.group(1))
            text = m.group(2).strip()
            htext = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            word_level = max(1, md_level - 1)  # ## -> 1, ### -> 2, #### -> 3
            doc.add_heading(htext, level=word_level)
            idx += 1
            continue

        # Horizontal rule
        if stripped == "---":
            _add_hr(doc)
            idx += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            _add_blockquote(doc, stripped[2:])
            idx += 1
            continue

        # Bullet list
        bm = re.match(r"^(\s*)-\s+(.*)$", line)
        if bm:
            indent = len(bm.group(1))
            level = 1 if indent >= 2 else 0
            _add_bullet(doc, bm.group(2), level=level)
            idx += 1
            continue

        # Blank line
        if stripped == "":
            idx += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        _add_runs_with_inline(p, stripped)
        idx += 1

    # ── Footer: page numbers + doc tag on the body section ──
    last_section = doc.sections[-1]
    _add_page_number_footer(last_section)
    _add_footer_text(last_section, footer_text)

    doc.save(docx_path)
    return docx_path


if __name__ == "__main__":
    out1 = convert(
        os.path.join(HERE, "PROJECT_REPORT.md"),
        os.path.join(HERE, "PROJECT_REPORT.docx"),
        cover_subtitle="Project Report",
        footer_text="OEMS — Online Examination & Monitoring System · Project Report",
    )
    print(f"[OK] {out1}")

    out2 = convert(
        os.path.join(HERE, "VIVA_READY_QA.md"),
        os.path.join(HERE, "VIVA_READY_QA.docx"),
        cover_subtitle="Viva & Interview Question Bank",
        footer_text="OEMS — Online Examination & Monitoring System · Viva Q&A",
    )
    print(f"[OK] {out2}")
